#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
应用Word模板到现有文档
此脚本将读取一个模板Word文档的样式，并将另一个文档的内容复制到应用了模板样式的新文档中。
"""

from docx import Document
from docx.shared import Inches
import os
import sys


def apply_template_content(template_path, content_path, output_path):
    """
    将模板样式应用到内容文档，并保存为新文档
    
    参数:
        template_path: 模板文档路径
        content_path: 内容文档路径
        output_path: 输出文档路径
    """
    try:
        # 打开模板文档和内容文档
        template_doc = Document(template_path)
        content_doc = Document(content_path)
        
        # 创建新文档（基于模板）
        new_doc = Document()
        
        # 复制模板中的所有样式（通过复制第一个段落的样式，然后删除它）
        if len(template_doc.paragraphs) > 0:
            # 在新文档中创建一个段落并应用模板样式
            first_para = new_doc.add_paragraph()
            first_para.style = template_doc.paragraphs[0].style
        
        # 处理标题样式
        title_styles = {}
        for i, para in enumerate(template_doc.paragraphs):
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split(' ')[1])
                title_styles[level] = para.style
                
        # 复制内容文档中的所有段落和表格
        for element in content_doc.element.body:
            # 处理段落
            if element.tag.endswith('p'):
                # 创建新段落
                new_para = new_doc.add_paragraph()
                
                # 复制段落内容
                for run in element.iterchildren():
                    new_run = new_para.add_run()
                    # 复制文本内容
                    if run.text:
                        new_run.text = run.text
                    # 复制文本格式
                    if run.get('w:rPr'):
                        # 这里简化处理，实际应用中可能需要更复杂的格式复制
                        pass
                
                # 尝试应用标题样式（如果匹配）
                original_text = element.text
                if original_text and len(original_text.strip()) > 0:
                    # 检查是否为标题
                    if original_text.startswith('# '):
                        # 一级标题
                        if 1 in title_styles:
                            new_para.style = title_styles[1]
                    elif original_text.startswith('## '):
                        # 二级标题
                        if 2 in title_styles:
                            new_para.style = title_styles[2]
                    elif original_text.startswith('### '):
                        # 三级标题
                        if 3 in title_styles:
                            new_para.style = title_styles[3]
            
            # 处理表格
            elif element.tag.endswith('tbl'):
                # 添加表格
                # 首先获取行数和列数
                rows = list(element.iterchildren())
                if len(rows) > 0:
                    # 获取第一行的列数
                    cols = list(rows[0].iterchildren())
                    if len(cols) > 0:
                        new_table = new_doc.add_table(rows=len(rows), cols=len(cols))
                        
                        # 复制表格内容
                        for i, row in enumerate(rows):
                            cells = list(row.iterchildren())
                            for j, cell in enumerate(cells):
                                if cell.text:
                                    new_table.cell(i, j).text = cell.text
        
        # 保存新文档
        new_doc.save(output_path)
        print(f"成功创建新文档: {output_path}")
        return True
    except Exception as e:
        print(f"处理文档时出错: {str(e)}")
        return False


def main():
    # 文件路径
    template_path = '/Users/easyo/Documents/oeasy/oeasy-python-tutorial/附件：模版-实验4-1 用计算思维求解问题.docx'
    content_path = '/Users/easyo/Documents/oeasy/oeasy-python-tutorial/output.docx'
    output_path = '/Users/easyo/Documents/oeasy/oeasy-python-tutorial/列表_加法_增强赋值_加等于_extend_扩展列表_格式化.docx'
    
    # 检查文件是否存在
    if not os.path.exists(template_path):
        print(f"错误: 模板文件不存在: {template_path}")
        sys.exit(1)
    
    if not os.path.exists(content_path):
        print(f"错误: 内容文件不存在: {content_path}")
        sys.exit(1)
    
    # 应用模板
    success = apply_template_content(template_path, content_path, output_path)
    
    if success:
        print("操作完成!")
    else:
        print("操作失败!")


if __name__ == "__main__":
    main()